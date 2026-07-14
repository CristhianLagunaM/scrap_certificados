from docx import Document
from docx.shared import RGBColor
import pandas as pd
import zipfile
import os

class EstudiosPreviosGenerator:

    def _copy_run_format(self, source_run, target_run):
        source_font = source_run.font
        target_font = target_run.font

        target_font.size = source_font.size
        target_font.bold = source_font.bold
        target_font.italic = source_font.italic
        target_font.underline = source_font.underline
        target_font.name = source_font.name
        target_font.all_caps = source_font.all_caps
        target_font.small_caps = source_font.small_caps
        target_font.strike = source_font.strike
        target_font.double_strike = source_font.double_strike
        target_font.subscript = source_font.subscript
        target_font.superscript = source_font.superscript
        target_font.shadow = source_font.shadow
        target_font.outline = source_font.outline
        target_font.rtl = source_font.rtl
        target_font.imprint = source_font.imprint
        target_font.emboss = source_font.emboss
        target_font.complex_script = source_font.complex_script
        target_font.cs_bold = source_font.cs_bold
        target_font.cs_italic = source_font.cs_italic
        target_font.highlight_color = source_font.highlight_color
        target_font.color.rgb = RGBColor(0, 0, 0)

        if source_run.style is not None:
            target_run.style = source_run.style

    def _apply_body_run_format(self, run):
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.all_caps = False
        run.font.small_caps = False
        run.font.strike = False
        run.font.double_strike = False
        run.font.subscript = False
        run.font.superscript = False
        run.font.shadow = False
        run.font.outline = False
        run.font.imprint = False
        run.font.emboss = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _normalize_placeholder(self, text):
        return " ".join(str(text).replace("\n", " ").split()).strip()

    def _replace_paragraph_text(self, paragraph, value):
        if not paragraph.runs:
            run = paragraph.add_run(value)
            self._apply_body_run_format(run)
            return

        first_run = paragraph.runs[0]
        first_run.text = value
        self._apply_body_run_format(first_run)

        for run in paragraph.runs[1:]:
            run.text = ""

    def replace_placeholder(self, container, search_variants, value):
        for paragraph in container.paragraphs:
            paragraph_text = self._normalize_placeholder(paragraph.text)
            exact_match = any(
                paragraph_text == self._normalize_placeholder(variant)
                for variant in search_variants
            )

            if exact_match:
                self._replace_paragraph_text(paragraph, value)
                continue

            for run in paragraph.runs:
                replacement_done = False
                for search_text in search_variants:
                    if self._normalize_placeholder(run.text) != self._normalize_placeholder(search_text):
                        continue

                    original_text = run.text
                    run.text = original_text.replace(original_text, value)
                    self._copy_run_format(run, run)
                    replacement_done = True

                if replacement_done:
                    continue

        if hasattr(container, "tables"):
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.replace_placeholder(cell, search_variants, value)

    def _placeholder_variants(self, column_name):
        clean_name = self._normalize_placeholder(column_name)
        variants = {
            clean_name,
            f"-{clean_name}",
            f"{{{{{clean_name}}}}}",
            f"<<{clean_name}>>",
        }
        return tuple(variant for variant in variants if variant)

    def generate(self, excel_path, word_path, output_folder):
        df = pd.read_excel(excel_path)

        zip_name = "estudios_previos.zip"
        zip_path = os.path.join(output_folder, zip_name)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for i, (_, fila) in enumerate(df.iterrows(), start=1):
                doc = Document(word_path)

                data = {
                    col: "" if pd.isna(fila[col]) else str(fila[col]).strip()
                    for col in df.columns
                }

                for placeholder, value in data.items():
                    self.replace_placeholder(
                        doc,
                        self._placeholder_variants(placeholder),
                        value,
                    )

                out_name = f"Estudios_Previos_{i}.docx"
                out_path = os.path.join(output_folder, out_name)
                doc.save(out_path)

                zipf.write(out_path, out_name)

        return zip_name
